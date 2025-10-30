#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script to merge two markdown reports and convert to Word document
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def add_heading(doc, text, level):
    """Add a heading to the document"""
    heading = doc.add_heading(text, level=level)
    heading_format = heading.paragraph_format
    heading_format.space_before = Pt(12)
    heading_format.space_after = Pt(6)
    return heading

def add_paragraph(doc, text, style=None):
    """Add a paragraph to the document"""
    if style:
        para = doc.add_paragraph(text, style=style)
    else:
        para = doc.add_paragraph(text)
    para_format = para.paragraph_format
    para_format.space_after = Pt(6)
    return para

def add_code_block(doc, code_text):
    """Add a code block to the document"""
    para = doc.add_paragraph(code_text)
    para.style = 'No Spacing'
    
    # Format as code
    for run in para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Add shading
    para_format = para.paragraph_format
    para_format.left_indent = Inches(0.5)
    para_format.space_before = Pt(6)
    para_format.space_after = Pt(6)
    
    return para

def parse_markdown_to_word(markdown_content, doc):
    """Parse markdown content and add to Word document"""
    lines = markdown_content.split('\n')
    i = 0
    in_code_block = False
    code_block_content = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_block_content)
                add_code_block(doc, code_text)
                code_block_content = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            add_heading(doc, line[2:].strip(), 1)
        elif line.startswith('## '):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith('### '):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith('#### '):
            add_heading(doc, line[5:].strip(), 4)
        elif line.startswith('##### '):
            add_heading(doc, line[6:].strip(), 5)
        
        # Handle horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph()
        
        # Handle bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            add_paragraph(doc, text, style='List Bullet')
        
        # Handle numbered lists
        elif re.match(r'^\d+\. ', line.strip()):
            text = re.sub(r'^\d+\. ', '', line.strip())
            add_paragraph(doc, text, style='List Number')
        
        # Handle blockquotes
        elif line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            para = add_paragraph(doc, text)
            para.paragraph_format.left_indent = Inches(0.5)
        
        # Handle regular paragraphs
        elif line.strip():
            # Skip lines that are markdown formatting
            if not line.strip().startswith('[') and not line.strip().startswith('**'):
                add_paragraph(doc, line.strip())
        
        i += 1

def create_word_document():
    """Main function to create Word document from markdown files"""
    
    # File paths
    script_dir = os.path.dirname(os.path.abspath(__file__))
    report1_path = os.path.join(script_dir, 'Report-1-Technical-Documentation.md')
    report2_path = os.path.join(script_dir, 'Report-2-System-Pipelines-and-Flows.md')
    output_path = os.path.join(script_dir, 'John-Henry-Complete-Technical-Report.docx')
    
    # Create new Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add title page
    title = doc.add_heading('BÁO CÁO KỸ THUẬT TOÀN DIỆN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('John Henry Fashion E-Commerce Platform')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(16)
    subtitle_format.bold = True
    
    doc.add_paragraph()
    
    info = doc.add_paragraph('Tác giả: AI Technical Documentation System')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info = doc.add_paragraph('Ngày tạo: 24/10/2025')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info = doc.add_paragraph('Phiên bản: 1.0')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Read Report 1
    print("Đang đọc Report 1...")
    with open(report1_path, 'r', encoding='utf-8') as f:
        report1_content = f.read()
    
    # Read Report 2
    print("Đang đọc Report 2...")
    with open(report2_path, 'r', encoding='utf-8') as f:
        report2_content = f.read()
    
    # Add Report 1 content
    print("Đang chuyển đổi Report 1...")
    heading = doc.add_heading('PHẦN A: TÀI LIỆU KỸ THUẬT', 1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    parse_markdown_to_word(report1_content, doc)
    
    doc.add_page_break()
    
    # Add Report 2 content
    print("Đang chuyển đổi Report 2...")
    heading = doc.add_heading('PHẦN B: SYSTEM PIPELINES VÀ FLOWS', 1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    parse_markdown_to_word(report2_content, doc)
    
    # Save document
    print(f"Đang lưu file Word...")
    doc.save(output_path)
    print(f"Hoàn thành! File đã được lưu tại: {output_path}")
    
    return output_path

if __name__ == '__main__':
    try:
        output_file = create_word_document()
        print(f"\n{'='*60}")
        print(f"Đã tạo file Word thành công!")
        print(f"Đường dẫn: {output_file}")
        print(f"{'='*60}\n")
    except Exception as e:
        print(f"Lỗi: {str(e)}")
        import traceback
        traceback.print_exc()
